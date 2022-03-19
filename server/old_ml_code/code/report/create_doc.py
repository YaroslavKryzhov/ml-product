import numpy as np

from docx import Document
from docx.shared import Inches
from sklearn.metrics import confusion_matrix, classification_report
import re
import os
import json
from datetime import datetime
from Web_classification_constructor_backend.settings import MEDIA_ROOT


# with open(os.path.join(f"{MEDIA_ROOT}", 'user_all_params.json')) as json_file:
#     all_params = json.load(json_file)


def table_to_doc(doc, header, nrows, ncols, matrix):
    """
    Вспомогательная функция для отчета
    """
    h = doc.add_heading(text=header, level=2)
    if header != '':
        doc.add_paragraph()
    table = doc.add_table(rows=nrows, cols=ncols)
    table.style = 'Table Grid'

    for row_num in range(nrows):
        row = table.rows[row_num]
        df_row = matrix[row_num]
        for i, (txt, cell) in enumerate(zip(df_row, row.cells)):
            cell.text = f'{txt}'
            if row_num == 0 or i == 0:
                run = cell.paragraphs[0].runs[0]
                run.font.bold = True


def create_heading(doc, text, level, alignment):
    """
    Вспомогательная функция для отчета
    """
    h = doc.add_heading(text=text, level=level)
    h.alignment = alignment
    doc.add_paragraph()


def add_graph(doc, img_path, imp_text, width, height):
    """
    Вспомогательная функция для отчета
    """
    p = doc.add_paragraph()
    p.alignment = 1
    r = p.add_run()
    r.add_picture(img_path, width=Inches(width), height=Inches(height))
    r.add_text(imp_text)


def get_confusion_matrix(estimators_pred, y_true):
    """
    Вспомогательная функция для отчета
    """
    confusion_m = {}
    for key, value in estimators_pred.items():

        tn, fp, fn, tp = confusion_matrix(y_true, value).ravel()
        matrix = ['', 'Prediction = 1', 'Prediction = 0', 'Sum',
                  'True = 1', tp, fn, tp + fn,
                  'True = 0', fp, tn, fp + tn,
                  'Sum', tp + fp, fn + tn, tp + fp + fn + tn]

        with open(os.path.join(f"{MEDIA_ROOT}",
                               'user_all_params.json')) as json_file:
            all_params = json.load(json_file)

        if all_params['common params']['composition method'] in ['voting', 'stacking']:
            confusion_m[key] = np.array(matrix).reshape(4, 4)
        else:
            for k, v in all_params['common params'].items():
                if k.find('number') != -1 and v != 0:
                    algorithm = k[:-7]  # костыль
            confusion_m[f'{algorithm} #1'] = np.array(matrix).reshape(4, 4)
    return confusion_m


def get_clssif_report(estimators_pred, y_true):
    """
    Вспомогательная функция для отчета
    """
    clssif_report = {}
    for key, value in estimators_pred.items():
        s = classification_report(y_true, value)
        s = re.sub('micro avg', 'micro_avg', s)
        s = re.sub('macro avg', 'macro_avg', s)
        s = re.sub('weighted avg', 'weighted_avg', s)
        s = re.sub('\n', ' ', s)
        report = [''] + [x for x in s.split(' ') if x != ''][:15] + ['', ''] + [x for x in s.split(' ') if x != ''][15:]
        with open(os.path.join(f"{MEDIA_ROOT}",
                               'user_all_params.json')) as json_file:
            all_params = json.load(json_file)
        if all_params['common params']['composition method'] in ['voting', 'stacking']:
            clssif_report[key] = np.array(report).reshape(6, 5)
        else:
            for k, v in all_params['common params'].items():
                if k.find('number') != -1 and v != 0:
                    algorithm = k[:-7]  # костыль
                    clssif_report[f'{algorithm} #1'] = np.array(report).reshape(6, 5)
    return clssif_report


def get_doc(pickle_path, train_path, feats_descr, pred_test, estimators_pred, y_true):
    """
    Главная  функция для отчета
    :param pickle_path: путь до модели
    :param train_path: путь до трейна
    :param feats_descr: описнаие фичей (глобальный объект)
    :param pred_test: предсказанные данные валидационной выборки
    :param estimators_pred:
    :param y_true: реальные данные валидационной выборки
    :return: None
    """
    doc = Document()
    width = 6
    height = 5

    create_heading(doc=doc, text=f'Отчет о результатах обучения модели от {datetime.today().strftime("%Y-%m-%d-%H-%M-%S")}',
                   level=0, alignment=0)
    create_heading(doc=doc, text='Характеристики обучающей выборки', level=1, alignment=1)

    table_to_doc(doc=doc, header='Описание признаков', nrows=feats_descr.shape[0], ncols=feats_descr.shape[1],
                 matrix=feats_descr)

    # Характеристики модели
    create_heading(doc=doc, text='Характеристики модели', level=1, alignment=1)

    # Model description
    # if all_params['default'] == 1:
    #     model_type = 'По умолчанию'
    # else:
    #     model_type = 'Собственная'
    with open(os.path.join(f"{MEDIA_ROOT}",
                           'user_all_params.json')) as json_file:
        all_params = json.load(json_file)

    composition_method = all_params['common params']['composition method']
    n_base_models = len(all_params['base algorithms'].keys())
    fs_method = all_params['common params']['feature selection method']
    fill_null_method = all_params['common params']['filling gaps method']
    anomalies_method = all_params['common params']['deleting anomalies method']

    confusion_m = get_confusion_matrix(estimators_pred, y_true)
    clssif_report = get_clssif_report(estimators_pred, y_true)

    doc.add_paragraph().add_run().add_text(f'Метод отбора признаков: {fs_method}')
    doc.add_paragraph().add_run().add_text(f'Метод заполнения пропусков: {fill_null_method}')
    doc.add_paragraph().add_run().add_text(f'Метод выявления аномалий: {anomalies_method}')
    # doc.add_paragraph().add_run().add_text(f'Тип модели: {model_type}')
    doc.add_paragraph().add_run().add_text(f'Метод композиции: {composition_method}')
    doc.add_paragraph().add_run().add_text(f'Количество базовых моделей: {n_base_models}')
    doc.add_paragraph().add_run().add_text(f'Базовые модели:')

    img = 1
    for i, key in enumerate(all_params['base algorithms'].keys()):
        doc.add_paragraph(f'{key}', style='List Bullet')
        doc.add_paragraph('Параметры:', style='List Bullet 2')
        for k, v in all_params['base algorithms'][key].items():
            doc.add_paragraph(f'{k}: {v}', style='List Bullet 3')
        doc.add_paragraph('Таблица сопряженности', style='List Bullet 2')
        table_to_doc(doc=doc, header='', nrows=4, ncols=4, matrix=confusion_m[key])
        doc.add_paragraph()
        doc.add_paragraph('Результаты классификации', style='List Bullet 2')
        table_to_doc(doc=doc, header='', nrows=6, ncols=5, matrix=clssif_report[key])
        doc.add_paragraph()
        doc.add_paragraph('Roc-кривая', style='List Bullet 2')

        if all_params['common params']['composition method'] in ['voting', 'stacking']:
            add_graph(doc=doc,
                      img_path=os.path.join(f"{MEDIA_ROOT}",
                                            'App/images/ROC_curve {}.png'.format(key)),
                      imp_text=f'Рис.{img} Roc-кривая для модели {key}', width=width, height=height)
            img += 1
            doc.add_paragraph()
            doc.add_paragraph('PR-кривая', style='List Bullet 2')
            add_graph(doc=doc,
                      img_path=os.path.join(f"{MEDIA_ROOT}",
                                            'App/images/PR_curve {}.png'.format(key)),
                      imp_text=f'Рис.{img} PR-кривая для модели {key}', width=width, height=height)
            img += 1
            doc.add_paragraph()
            doc.add_paragraph('Distribution graph', style='List Bullet 2')
            add_graph(doc=doc,
                      img_path=os.path.join(f"{MEDIA_ROOT}",
                                            'App/images/Distribution_graph {}.png'.format(key)),
                      imp_text=f'Рис.{img} Гистограмма распределения предсказаний модели {key}', width=width,
                      height=height)
            img += 1
            doc.add_paragraph()

        else:
            pass

    # Характеристики прогноза по итоговой модели
    if all_params['common params']['composition method'] in ['voting', 'stacking']:
        create_heading(doc, text='Характеристики прогноза по итоговой модели', level=1, alignment=1)
        table_to_doc(doc=doc, header='Таблица сопряженности', nrows=4, ncols=4, matrix=confusion_m['final_model'])
        table_to_doc(doc=doc, header='Результаты классификации', nrows=6, ncols=5, matrix=clssif_report['final_model'])

    # Визуализация результатов
    doc.add_paragraph()
    create_heading(doc, text='Визуализация результатов', level=2, alignment=0)

    add_graph(doc=doc,
              img_path=os.path.join(f"{MEDIA_ROOT}", 'App/images/ROC_curve.png'),
              imp_text=f'Рис.{img} Roc-кривая', width=width, height=height)
    img += 1
    add_graph(doc=doc,
              img_path=os.path.join(f"{MEDIA_ROOT}", 'App/images/PR_curve.png'),
              imp_text=f'Рис.{img} PR-кривая', width=width, height=height)
    img += 1
    add_graph(doc=doc,
              img_path=os.path.join(f"{MEDIA_ROOT}", 'App/images/PR_by_prc.png'),
              imp_text=f'Рис.{img} График precision и recall по перцентилям', width=width, height=height)
    img += 1
    add_graph(doc=doc, img_path=os.path.join(f"{MEDIA_ROOT}",
                                             'App/images/Distribution_graph final_model.png'),
              imp_text=f'Рис.{img} Гистограмма распределения предсказаний итоговой модели', width=width, height=height)

    test_ex = np.vstack((pred_test.columns, pred_test.head(10).values))
    table_to_doc(doc=doc, header='Предсказание на тестовой выборке', nrows=11, ncols=3, matrix=test_ex)

    path_to_final_word = os.path.join(f"{MEDIA_ROOT}",
                                      'Output/fin.docx')
    doc.save(path_to_final_word)
